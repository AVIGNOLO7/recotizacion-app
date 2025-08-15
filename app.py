import streamlit as st
import pandas as pd
import numpy as np
import re
import io
from datetime import datetime
from docx import Document

def norm_money_ar(s: str):
    s = str(s).replace('$','').replace(' ','')
    s = s.replace('.','').replace(',', '.')
    try:
        return float(s)
    except:
        return np.nan

def extraer_franquicia_texto(cobertura: str):
    m = re.search(r'(\d{1,2})\s*%', str(cobertura))
    return float(m.group(1)) if m else np.nan

def moneda(x):
    try:
        return f"${x:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except:
        return ""

def parse_txt_options(texto: str):
    rows = []
    blob = texto.replace('\n', ' ')
    # Zurich: DV/D1/D2/D6 -> 4%/1%/2%/6%
    Z_MAP = {'DV':4, 'D1':1, 'D2':2, 'D6':6}
    if re.search(r'ZURICH', blob, flags=re.I):
        for code, fr in Z_MAP.items():
            m = re.search(rf'{code}[^0-9]*(\d{{1,3}}(?:\.\d{{3}})*,\d{{2}})', blob)
            if m:
                rows.append({'Compañía':'Zurich','Código':code,'Cobertura':'Todo riesgo','Franquicia (%)':fr,'Prima mensual ($)':norm_money_ar(m.group(1))})
    # Allianz: D 1% / D 2% / D 3% / D 5% / D 6%
    if re.search(r'ALLIANZ', blob, flags=re.I):
        for fr in [1,2,3,4,5,6]:
            m = re.search(rf'D\s*{fr}%[^0-9]*(\d{{1,3}}(?:\.\d{{3}})*,\d{{2}})', blob)
            if m:
                rows.append({'Compañía':'Allianz','Código':f'D {fr}%','Cobertura':'Todo riesgo','Franquicia (%)':fr,'Prima mensual ($)':norm_money_ar(m.group(1))})
    # Sura: TR 1% / TR 2% / TR 3% / TR 4% / ...
    if re.search(r'SURA', blob, flags=re.I):
        for fr in [1,2,3,4,5,6]:
            m = re.search(rf'TR\s*{fr}%[^0-9]*(\d{{1,3}}(?:\.\d{{3}})*,\d{{2}})', blob)
            if m:
                rows.append({'Compañía':'Sura','Código':f'TR {fr}%','Cobertura':'Todo riesgo','Franquicia (%)':fr,'Prima mensual ($)':norm_money_ar(m.group(1))})
    df = pd.DataFrame(rows).drop_duplicates()
    if not df.empty:
        df = df.sort_values(['Franquicia (%)','Compañía','Prima mensual ($)'])
    return df

def aplicar_reglas(df, vigente):
    comp_act = vigente['comp_act']
    cob_act = vigente['cob_act']
    franq_act = vigente['franq_act']
    prima_act = vigente['prima_act']
    objetivo = vigente['objetivo']

    if df is None or df.empty:
        return df, pd.DataFrame(), pd.DataFrame(), pd.DataFrame()

    df = df.copy()
    df['Elegible'] = df['Prima mensual ($)'] <= prima_act

    def is_superior(row):
        try:
            return float(row['Franquicia (%)']) < float(franq_act)
        except:
            return False

    descartar_superior = (objetivo.lower() != 'mejorar cobertura')
    df['Superior'] = df.apply(is_superior, axis=1) if descartar_superior else False
    df['Incluida 10%'] = df['Elegible'] & (~df['Superior'])

    df['Ahorro ($)'] = (prima_act - df['Prima mensual ($)']).round(2)
    df['Ahorro (%)'] = ((1 - (df['Prima mensual ($)']/prima_act))*100).round(1)

    # Mejor equivalente y mejor inferior (si existen)
    eq = df[(df['Franquicia (%)']==franq_act) & (df['Incluida 10%'])].sort_values('Prima mensual ($)').head(1)
    inf = df[(df['Franquicia (%)']>franq_act) & (df['Incluida 10%'])].sort_values('Prima mensual ($)').head(1)

    # Regla: NO aumentar comisión si es misma compañía y misma cobertura (no se aplica acá)
    presentables = pd.concat([eq, inf]).drop_duplicates()

    return df, eq, inf, presentables

def generar_informe_interno(df, vigente):
    doc = Document()
    doc.add_heading('Informe interno – Recotización', level=1)
    doc.add_paragraph(f"Cliente: {vigente['cliente']}")
    doc.add_paragraph(f"Vigente: {vigente['comp_act']} – {vigente['cob_act']} – Franq {vigente['franq_act']:.0f}% – Prima {moneda(vigente['prima_act'])} – Comisión {vigente['margen_act']:.0f}%")
    doc.add_paragraph(f"Objetivo: {vigente['objetivo']}")
    doc.add_paragraph('')

    doc.add_heading('Tabla comparativa (10%)', level=2)
    tbl = doc.add_table(rows=1, cols=6)
    hdr = tbl.rows[0].cells
    hdr[0].text='Franquicia'; hdr[1].text='Compañía'; hdr[2].text='Prima'; hdr[3].text='Elegible'; hdr[4].text='Superior?'; hdr[5].text='Ahorro %'
    for _, r in df.sort_values(['Franquicia (%)','Prima mensual ($)']).iterrows():
        row = tbl.add_row().cells
        row[0].text = f"{int(r['Franquicia (%)'])}%"
        row[1].text = str(r['Compañía'])
        row[2].text = moneda(r['Prima mensual ($)'])
        row[3].text = 'Sí' if r['Elegible'] else 'No'
        row[4].text = 'Sí' if r['Superior'] else 'No'
        row[5].text = f"{r['Ahorro (%)']:.1f}%"

    doc.add_paragraph('')
    doc.add_heading('Conclusión', level=2)
    if df['Incluida 10%'].sum() == 0:
        doc.add_paragraph('El cliente ya se encuentra en la mejor alternativa; no hay opciones más económicas para presentar. No corresponde aumento de comisión en misma compañía/cobertura.')
    else:
        doc.add_paragraph('Se identificaron opciones presentables según reglas (ver Informe Cliente).')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generar_informe_cliente(eq, inf, vigente):
    doc = Document()
    doc.add_heading('Propuesta de renovación', level=1)
    doc.add_paragraph(f"Estimado/a {vigente['cliente']},")
    doc.add_paragraph(f"Analizamos su póliza vigente ({vigente['cob_act']}, franquicia {vigente['franq_act']:.0f}%) y comparamos con las principales aseguradoras.")

    if (eq is None or eq.empty) and (inf is None or inf.empty):
        doc.add_paragraph('Confirmamos que su cobertura actual se mantiene como la alternativa más conveniente en su categoría. No es necesario realizar cambios para mejorar el costo.')
        doc.add_paragraph(f"Compañía: {vigente['comp_act']}")
        doc.add_paragraph(f"Cobertura: {vigente['cob_act']}")
        doc.add_paragraph(f"Prima mensual: {moneda(vigente['prima_act'])}")
    else:
        if eq is not None and not eq.empty:
            r = eq.iloc[0]
            doc.add_paragraph('Opción equivalente recomendada:')
            doc.add_paragraph(f"• {r['Compañía']} – Franquicia {int(r['Franquicia (%)'])}% – {moneda(r['Prima mensual ($)'])} (ahorro {r['Ahorro (%)']}%)")
        if inf is not None and not inf.empty:
            r = inf.iloc[0]
            doc.add_paragraph('Opción alternativa (franquicia mayor):')
            doc.add_paragraph(f"• {r['Compañía']} – Franquicia {int(r['Franquicia (%)'])}% – {moneda(r['Prima mensual ($)'])} (ahorro {r['Ahorro (%)']}%)")

    doc.add_paragraph('')
    doc.add_paragraph('Quedamos a disposición por cualquier aclaración que estime conveniente.')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_excel(df, vigente):
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine='openpyxl') as w:
        pd.DataFrame({
            'Campo':['Cliente','Compañía actual','Cobertura actual','Franquicia actual (%)','Prima actual ($)','Margen actual (%)','Objetivo'],
            'Valor':[vigente['cliente'], vigente['comp_act'], vigente['cob_act'], vigente['franq_act'], vigente['prima_act'], vigente['margen_act'], vigente['objetivo']]
        }).to_excel(w, sheet_name='Vigente', index=False)
        df.to_excel(w, sheet_name='Comparativa 10%', index=False)
        df[df['Incluida 10%']].to_excel(w, sheet_name='Filtrado', index=False)
    out.seek(0)
    return out

st.title("Recotización de Autos – San Jorge")
st.markdown("Subí el **Excel Vigente** y pegá el **texto** de la tabla del PDF 10%.")

xlsx = st.file_uploader("Excel Vigente (.xlsx)", type=['xlsx'])
pdf_txt = st.text_area("Texto del PDF (pegado)", height=200)

if st.button("Analizar"):
    if not xlsx or not pdf_txt:
        st.error("Cargá el Excel y pegá el texto del PDF para continuar.")
    else:
        try:
            vigdf = pd.read_excel(xlsx, sheet_name='Vigente').iloc[0]
        except Exception as e:
            st.error(f"No pude leer la hoja 'Vigente' del Excel: {e}")
            st.stop()

        cliente = str(vigdf['Cliente']).strip()
        comp_act = str(vigdf['Compañía actual']).strip()
        cob_act = str(vigdf['Cobertura actual']).strip()
        franq_act = vigdf.get('Franquicia actual (%)', np.nan)
        if pd.isna(franq_act):
            franq_act = extraer_franquicia_texto(cob_act)
        prima_act = float(vigdf['Prima actual ($)'])
        margen_act = float(vigdf['Margen actual (%)'])
        objetivo = str(vigdf['Objetivo']).strip()

        vigente = dict(cliente=cliente, comp_act=comp_act, cob_act=cob_act, franq_act=float(franq_act),
                       prima_act=prima_act, margen_act=margen_act, objetivo=objetivo)

        df = parse_txt_options(pdf_txt)
        if df is None or df.empty:
            st.error("No pude reconocer opciones en el texto pegado. Revisá el formato (DV/D1/D2, D 1%, TR 1%, etc.).")
            st.stop()

        df, eq, inf, presentables = aplicar_reglas(df, vigente)

        st.subheader("Tabla comparativa (10%)")
        st.dataframe(df)

        buf_int = generar_informe_interno(df, vigente)
        buf_cli = generar_informe_cliente(eq, inf, vigente)
        xlsx_buf = export_excel(df, vigente)

        st.download_button("Descargar Informe Interno (DOCX)", data=buf_int, file_name=f"Informe_Interno_{cliente}.docx")
        st.download_button("Descargar Informe Cliente (DOCX)", data=buf_cli, file_name=f"Informe_Cliente_{cliente}.docx")
        st.download_button("Descargar Excel de respaldo (XLSX)", data=xlsx_buf, file_name=f"Recotizacion_{cliente}.xlsx")
        st.success("Listo. Descargá los archivos generados.")
